"""Bulk file loader: scans a folder for supported document types, reads text
with automatic encoding detection, and returns (doc_id, text, source) tuples
for indexing. Reports per-file errors rather than failing the whole batch."""

from __future__ import annotations

from pathlib import Path
from typing import Any

SUPPORTED_SUFFIXES = {".txt", ".md", ".docx", ".rtf"}


def _detect_encoding(raw: bytes) -> str:
    """Try to decode `raw` with charset-normalizer; fall back to utf-8 with
    surrogate escaping so we never lose data for the user."""
    try:
        from charset_normalizer import from_bytes
    except ImportError:
        return raw.decode("utf-8", errors="replace")

    result = from_bytes(raw).best()
    if result is None:
        return raw.decode("utf-8", errors="replace")
    return str(result)


def _read_txt(path: Path) -> str:
    raw = path.read_bytes()
    return _detect_encoding(raw)


def _read_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def _read_rtf(path: Path) -> str:
    from striprtf.striprtf import rtf_to_text

    raw = path.read_text(encoding="utf-8", errors="replace")
    return rtf_to_text(raw)


_READERS = {
    ".txt": _read_txt,
    ".md": _read_txt,
    ".docx": _read_docx,
    ".rtf": _read_rtf,
}


def load_folder(
    folder: str | Path,
    source_label: str = "folder",
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    """Walk *folder* recursively, read every supported file, return
    `(successes, errors)` where each success is ``{"doc_id": path_stem,
    "text": content, "source": source_label}`` and each error is
    ``{"path": …, "error": message}``."""
    root = Path(folder).expanduser().resolve()
    if not root.is_dir():
        return [], [{"path": str(root), "error": "Not a directory"}]

    successes: list[dict[str, Any]] = []
    errors: list[dict[str, Any]] = []

    for file_path in sorted(root.rglob("*")):
        if not file_path.is_file():
            continue
        suffix = file_path.suffix.lower()
        if suffix not in _READERS:
            continue

        reader = _READERS[suffix]
        try:
            text = reader(file_path)
        except Exception as exc:
            errors.append(
                {
                    "path": str(file_path),
                    "suffix": suffix,
                    "error": f"{type(exc).__name__}: {exc}",
                }
            )
            continue

        if not text or not text.strip():
            errors.append(
                {"path": str(file_path), "suffix": suffix, "error": "Empty or unreadable content"}
            )
            continue

        successes.append(
            {
                "doc_id": file_path.stem,
                "text": text.strip(),
                "source": source_label,
            }
        )

    return successes, errors


def load_files(
    files: list[tuple[str, bytes, str]],
    source_label: str = "upload",
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    """Process a list of ``(filename, raw_bytes, suffix)`` uploaded from the
    browser (where we only have the file *contents*, not a disk path).

    Returns ``(successes, errors)`` in the same shape as :func:`load_folder`
    so the indexing pipeline is identical from this point onward.
    """
    successes: list[dict[str, Any]] = []
    errors: list[dict[str, Any]] = []

    for filename, raw_bytes, suffix in files:
        suffix_lower = suffix.lower()
        if suffix_lower not in _READERS:
            errors.append(
                {"path": filename, "suffix": suffix, "error": f"Unsupported file type: {suffix}"}
            )
            continue

        reader = _READERS[suffix_lower]
        doc_id = Path(filename).stem

        try:
            if suffix_lower in (".txt", ".md"):
                text = _detect_encoding(raw_bytes)
            elif suffix_lower == ".docx":
                from io import BytesIO

                from docx import Document

                doc = Document(BytesIO(raw_bytes))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                text = "\n".join(paragraphs)
            elif suffix_lower == ".rtf":
                from striprtf.striprtf import rtf_to_text

                text = rtf_to_text(raw_bytes.decode("utf-8", errors="replace"))
            else:
                errors.append(
                    {"path": filename, "suffix": suffix, "error": f"No reader for {suffix}"}
                )
                continue
        except Exception as exc:
            errors.append(
                {
                    "path": filename,
                    "suffix": suffix,
                    "error": f"{type(exc).__name__}: {exc}",
                }
            )
            continue

        if not text or not text.strip():
            errors.append(
                {"path": filename, "suffix": suffix, "error": "Empty or unreadable content"}
            )
            continue

        successes.append(
            {
                "doc_id": doc_id,
                "text": text.strip(),
                "source": source_label,
            }
        )

    return successes, errors
